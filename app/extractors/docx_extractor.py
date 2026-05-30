"""DOCX document extraction using python-docx.

Extracts text from all paragraphs, table rows, headers, and footers
of a .docx file.
"""

from io import BytesIO
from typing import Optional

from ..extractor import ExtractionResult
from .base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extractor for .docx (Office Open XML) documents.

    Uses python-docx to extract text content from paragraphs, tables,
    headers, and footers. Returns an ExtractionResult with full_text,
    description (first meaningful paragraph), and confidence based on
    text length.
    """

    def extract(self, data: bytes, filename: Optional[str] = None) -> ExtractionResult:
        """Extract text from a .docx document.

        Args:
            data: Raw .docx file bytes.
            filename: Optional original filename.

        Returns:
            ExtractionResult with extracted text content.
        """
        try:
            from docx import Document as DocxDocument  # type: ignore[import-not-found]
            doc = DocxDocument(BytesIO(data))
        except Exception:
            return ExtractionResult(
                full_text="",
                confidence=0.0,
                requirements=["Failed to open DOCX document"],
            )

        all_text_parts: list[str] = []
        paragraphs: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                all_text_parts.append(text)

        # Extract tables (row by row, tab-separated)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = "\t".join(c for c in cells if c)
                if row_text:
                    all_text_parts.append(row_text)

        # Extract headers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_text_parts.append(text)

        # Extract footers
        for section in doc.sections:
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_text_parts.append(text)

        full_text = "\n".join(all_text_parts)

        if not full_text.strip():
            return ExtractionResult(
                full_text="",
                confidence=0.0,
                requirements=["No extractable text found in DOCX document"],
            )

        # Description: first meaningful paragraph (length > 50 chars)
        description = ""
        for para in paragraphs:
            if len(para) > 50:
                description = para
                break

        # Confidence based on text length
        char_count = len(full_text)
        if char_count >= 500:
            confidence = 0.7
        elif char_count >= 100:
            confidence = 0.5
        else:
            confidence = 0.3

        return ExtractionResult(
            full_text=full_text,
            description=description,
            confidence=confidence,
            raw_text_preview=full_text[:500] if full_text else None,
        )
